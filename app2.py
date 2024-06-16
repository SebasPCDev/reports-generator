
from docx import Document
from datetime import datetime
import os

def replace_text_in_paragraph(paragraph, old_text, new_text):
    if old_text in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if old_text in item.text:
                item.text = item.text.replace(old_text, new_text)

def replace_text_in_table(table, old_text, new_text):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_text_in_paragraph(paragraph, old_text, new_text)

def create_copies(template_path, output_dir, serial_numbers, places):
    
    # Ensure the output directory exists
    os.makedirs(output_dir, exist_ok=True)
    
    for serial, location in zip(serial_numbers, places):
        # Create a copy of the template document
        new_doc = Document(template_path)
        
        

        print(f'Creating copy  with serial number: {serial}')
        
        # Replace the placeholders in the document
        # for paragraph in new_doc.paragraphs:
            # replace_text_in_paragraph(paragraph, 'PLACEHOLDER_DATE', new_date)
            # replace_text_in_paragraph(paragraph, 'PLACEHOLDER_SERIAL', new_serial)

        for table in new_doc.tables:
            replace_text_in_table(table, 'PLACEHOLDER_SERIAL', serial)
            replace_text_in_table(table, 'PLACEHOLDER_LOCATION', location)
        
        # Define the output file name
        output_filename = f'RMP 2024-06-16 {serial}.docx'
        output_path = os.path.join(output_dir, output_filename)
        
        # Save the new document
        new_doc.save(output_path)
        print(f'Created: {output_path}')

# Parameters
template_path = 'templatecamillas.docx'
output_dir = 'camillas'
serial_numbers = [
    "4737100018660039",
    "4788100018660012",
    "011506100022690009",
    "4788100018660035",
    "4737100018660038",
    "4788100018660002",
    "4737100018660033",
    "4788100018660001",
    "4737100018660042",
    "4788100018660053",
    "4788100018660066",
    "4788100018660011",
    "4932100018660011",
    "011505100024070032",
    "011505100024070019",
    "011505100024070020",
    "4788100018660013",
    "4788100018660047",
    "011505100024070034",
    "011505100024070024",
    "011506100022690073",
    "011505100024070023",
    "4737100018660041",
    "011506100022690024",
    "011506100022690015",
    "01135910002690041",
    "011505100024070022",
    "011506100022690072",
    "011506100022690041",
    "011506100022690057",
    "011506100022690073",
    "011505100024070023"
]


places = [
    "URGENCIAS",
    "URGENCIAS",
    "CONSULTORIO 3",
    "URGENCIAS",
    "URGENCIAS",
    "URGENCIAS",
    "URGENCIAS",
    "URGENCIAS",
    "URGENCIAS",
    "URGENCIAS",
    "URGENCIAS",
    "URGENCIAS",
    "URGENCIAS",
    "URGENCIAS",
    "URGENCIAS",
    "URGENCIAS",
    "ECOGRAFIA",
    "URGENCIAS",
    "URGENCIAS",
    "URGENCIAS",
    "URGENCIAS",
    "URGENCIAS",
    "CIRUGIA / RECUPERACION",
    "CIRUGIA / RECUPERACION",
    "CIRUGIA / RECUPERACION",
    "CIRUGIA / RECUPERACION",
    "CIRUGIA / RECUPERACION",
    "CIRUGIA / RECUPERACION",
    "CIRUGIA / RECUPERACION",
    "CIRUGIA / RECUPERACION",
    "CIRUGIA / RECUPERACION"
]



# Run the script    
create_copies(template_path, output_dir, serial_numbers, places)
