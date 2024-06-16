from docx import Document
from datetime import datetime
import os

def replace_text_in_paragraph(paragraph, old_text, new_text):
    if old_text in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if old_text in item.text:
                item.text = item.text.replace(old_text, new_text)

def replace_text_in_table(table, old_text, new_text):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_text_in_paragraph(paragraph, old_text, new_text)

def create_copies(template_path, output_dir, serial_numbers, rooms_names):
    
    # Ensure the output directory exists
    os.makedirs(output_dir, exist_ok=True)
    
    for serial, room in zip(serial_numbers, rooms_names):
        # Create a copy of the template document
        new_doc = Document(template_path)
        
        

        print(f'Creating copy  with serial number: {serial}')
        
        # Replace the placeholders in the document
        # for paragraph in new_doc.paragraphs:
            # replace_text_in_paragraph(paragraph, 'PLACEHOLDER_DATE', new_date)
            # replace_text_in_paragraph(paragraph, 'PLACEHOLDER_SERIAL', new_serial)

        for table in new_doc.tables:
            replace_text_in_table(table, 'PLACEHOLDER_SERIAL', serial)
            replace_text_in_table(table, 'PLACEHOLDER_LOCATION', room)
        
        # Define the output file name
        output_filename = f'RMP 2024-06-16 {serial} {room}.docx'
        output_path = os.path.join(output_dir, output_filename)
        
        # Save the new document
        new_doc.save(output_path)
        print(f'Created: {output_path}')

# Parameters
template_path = 'template.docx'
output_dir = 'output'
serial_numbers = [
    "5080100019410001",
    "5080100019410002",
    "5080100019410007",
    "5080100019410008",
    "5080100019410009",
    "5080100019410010",
    "5080100019410012",
    "5080100019410013",
    "5080100019410014",
    "008100100020290057",
    "008100100020290074",
    "008100100020290099",
    "008100100020290151",
    "008100100020290156",
    "5080100019410016",
    "680210002031002 1",
    "680210002031002 3",
    "680210002031002 9",
    "008895100020290219",
    "008895100020290086",
    "008895100020290089",
    "008100100020290261",
    "008895100020290174",
    "008895100020290199",
    "008895100020290048",
    "008895100020290244",
    "008895100020290231",
    "008895100020290171"
]

room_names = [
    "Clinica GPP - HABITACION 24",
    "Clinica GPP - HABITACION 19",
    "Clinica GPP - HABITACION 18",
    "Clinica GPP - HABITACION 23",
    "Clinica GPP - HABITACION 7",
    "Clinica GPP - HABITACION 16",
    "Clinica GPP - HABITACION 21",
    "Clinica GPP - HABITACION 11",
    "Clinica GPP - HABITACION 5",
    "Clinica GPP - HABITACION 8",
    "Clinica GPP - HABITACION 5",
    "Clinica GPP - HABITACION 22",
    "Clinica GPP - HABITACION 4",
    "Clinica GPP - HABITACION 6",
    "Clinica GPP - HABITACION 17",
    "Clinica GPP - HABITACION 27",
    "Clinica GPP - HABITACION 28",
    "Clinica GPP - HABITACION 26",
    "Clinica GPP - HABITACION 11",
    "Clinica GPP - HABITACION 20",
    "Clinica GPP - HABITACION 9",
    "Clinica GPP - HABITACION 2",
    "Clinica GPP - HABITACION 3",
    "Clinica GPP - HABITACION 15",
    "Clinica GPP - HABITACION 1",
    "Clinica GPP - HABITACION 4",
    "Clinica GPP - HABITACION 10",
    "Clinica GPP - HABITACION 12"
]


# Run the script    
create_copies(template_path, output_dir, serial_numbers, room_names)
